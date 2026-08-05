"""
SCADA Design Narrative — Word output.

Stage 3.  The LLM writes Markdown because that is what it is reliably good
at, and because a plain-text intermediate stays diffable and reviewable.
This module renders that Markdown into a Word document an engineer can hand
to a client.

It is a deliberately narrow Markdown renderer, not a general one: it handles
the constructs the narrative prompts actually produce — headings, tables,
fenced code, block quotes, lists and inline emphasis — and formats them for
a technical document.  Tables get the most attention because a design
narrative is mostly tables: shaded repeating header rows, banded body rows,
and hairline borders that stay legible when the point list runs three pages.

python-docx rather than pandoc: it is a pip dependency instead of a system
package, which matters because this runs in the sidecar container.
"""

import io
import re
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# ─── Palette ─────────────────────────────────────────────────────────────

INK = RGBColor(0x1A, 0x1A, 0x1A)
SLATE = RGBColor(0x33, 0x41, 0x55)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x60, 0x6A, 0x7B)

HEADER_FILL = "1F4E79"
BAND_FILL = "F2F5F8"
CODE_FILL = "F4F4F6"
RULE_COLOR = "B8C2CC"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

# ─── Markdown patterns ───────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_TABLE_DIVIDER_RE = re.compile(r"^\s*\|?[\s:\-|]+\|[\s:\-|]*$")
_FENCE_RE = re.compile(r"^\s*```")
_QUOTE_RE = re.compile(r"^>\s?(.*)$")
_ULIST_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_OLIST_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_RULE_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")

# Inline emphasis, longest marker first so ** wins over *.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))", re.S)


# ─── Low-level Word helpers ──────────────────────────────────────────────


def _shade(element, fill: str) -> None:
    """Apply a solid background fill to a cell or paragraph."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _set_cell_margins(table, top=60, bottom=60, left=110, right=110) -> None:
    """Breathing room inside cells, in twentieths of a point."""
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for name, value in (
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def _set_table_borders(table, color: str = RULE_COLOR, size: int = 4) -> None:
    """Hairline borders — heavy grid lines fight with dense numeric content."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    table._tbl.tblPr.append(borders)


def _repeat_header_row(row) -> None:
    """Mark a row as a header so Word repeats it across page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _keep_with_next(paragraph) -> None:
    """Stop a heading being orphaned at the foot of a page."""
    paragraph.paragraph_format.keep_with_next = True


def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field code (used for page numbers)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# ─── Inline formatting ───────────────────────────────────────────────────


def _add_inline(
    paragraph,
    text: str,
    base_size: Optional[Pt] = None,
    color=INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """
    Render inline Markdown emphasis into runs.

    Tag names and expressions arrive wrapped in backticks throughout the
    narrative, so inline code is the most important of these to get right —
    it is what makes `PQM7_DNP.W3.instMag` readable in running prose.

    Emphasis nests: the findings section writes things like
    ``**2. `PowinMode` is hard-coded.**``, so bold and italic recurse into
    their own content rather than emitting it as one literal run.
    """
    for part in _INLINE_RE.split(text):
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            _add_inline(paragraph, part[2:-2], base_size, color, True, italic)
            continue

        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            _add_inline(paragraph, part[1:-1], base_size, color, bold, True)
            continue

        if part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.color.rgb = SLATE
            run.font.size = Pt(base_size.pt - 0.5) if base_size else Pt(9.5)
        else:
            run = paragraph.add_run(part)
            if base_size:
                run.font.size = base_size
            run.font.color.rgb = color

        run.bold = bold
        run.italic = italic


# ─── Document scaffolding ────────────────────────────────────────────────


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    specs = {
        "Heading 1": (16, True, ACCENT, 18, 8),
        "Heading 2": (13, True, ACCENT, 14, 6),
        "Heading 3": (11.5, True, SLATE, 12, 4),
        "Heading 4": (10.5, True, SLATE, 10, 4),
    }
    for name, (size, bold, color, before, after) in specs.items():
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_footer(document: Document, label: str) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(f"{label}    |    Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.name = BODY_FONT

    _add_field(paragraph, " PAGE ")
    run = paragraph.add_run(" of ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _add_field(paragraph, " NUMPAGES ")

    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
        run.font.name = BODY_FONT


def _add_title_block(
    document: Document,
    title: str,
    subtitle: str,
    metadata: Optional[Dict[str, str]] = None,
) -> None:
    for _ in range(3):
        document.add_paragraph()

    paragraph = document.add_paragraph()
    run = paragraph.add_run("SCADA DESIGN NARRATIVE")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = MUTED
    run.font.name = BODY_FONT
    paragraph.paragraph_format.space_after = Pt(2)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.name = BODY_FONT
    paragraph.paragraph_format.space_after = Pt(4)

    if subtitle:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = SLATE
        run.font.name = BODY_FONT
        paragraph.paragraph_format.space_after = Pt(18)

    if metadata:
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_borders(table, color="DDE3EA", size=2)
        _set_cell_margins(table)
        for key, value in metadata.items():
            cells = table.add_row().cells
            cells[0].width = Inches(2.1)
            cells[1].width = Inches(4.2)

            paragraph = cells[0].paragraphs[0]
            run = paragraph.add_run(key)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
            run.font.name = BODY_FONT

            paragraph = cells[1].paragraphs[0]
            _add_inline(paragraph, value, base_size=Pt(9.5))

    document.add_page_break()


# ─── Block renderers ─────────────────────────────────────────────────────


def _split_row(line: str) -> List[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [c.strip() for c in stripped.split("|")]


def _column_widths(rows: List[List[str]], total: Inches) -> List[Emu]:
    """
    Distribute width by the longest cell in each column.

    Naive equal columns waste the page on tables that pair a short DNP index
    against a long source expression, which is most of them.  Widths are
    clamped so no single column collapses or swallows the table.
    """
    column_count = max(len(r) for r in rows)
    longest = [1] * column_count
    for row in rows:
        for i, cell in enumerate(row):
            longest[i] = max(longest[i], len(cell))

    # Compress the dynamic range — a cell 10x longer should not get 10x the
    # width, or short columns become unreadably narrow.
    weights = [max(1.0, length) ** 0.62 for length in longest]
    total_weight = sum(weights)

    min_share = 0.07
    shares = [max(min_share, w / total_weight) for w in weights]
    scale = sum(shares)
    return [Emu(int(total.emu * (s / scale))) for s in shares]


def _render_table(document: Document, block: List[str], content_width: Inches) -> None:
    rows = [_split_row(line) for line in block if not _TABLE_DIVIDER_RE.match(line)]
    if not rows:
        return

    column_count = max(len(r) for r in rows)
    rows = [r + [""] * (column_count - len(r)) for r in rows]
    header, body = rows[0], rows[1:]

    table = document.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_borders(table)
    _set_cell_margins(table)

    widths = _column_widths(rows, content_width)

    header_cells = table.rows[0].cells
    for i, text in enumerate(header):
        cell = header_cells[i]
        cell.width = widths[i]
        _shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(re.sub(r"[*`]", "", text))
        run.font.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = BODY_FONT

    _repeat_header_row(table.rows[0])

    for index, row_values in enumerate(body):
        cells = table.add_row().cells
        _prevent_row_split(table.rows[-1])
        for i, text in enumerate(row_values):
            cell = cells[i]
            cell.width = widths[i]
            if index % 2 == 1:
                _shade(cell._tc.get_or_add_tcPr(), BAND_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            _add_inline(paragraph, text, base_size=Pt(8.5))

    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _render_code(document: Document, lines: List[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    _shade(paragraph._p.get_or_add_pPr(), CODE_FILL)

    for i, line in enumerate(lines):
        if i:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = SLATE


def _render_quote(document: Document, lines: List[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    _shade(paragraph._p.get_or_add_pPr(), BAND_FILL)
    _add_inline(paragraph, " ".join(lines), base_size=Pt(9))
    for run in paragraph.runs:
        run.italic = True
        run.font.color.rgb = MUTED


# ─── Markdown walk ───────────────────────────────────────────────────────


def _is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def render_markdown(document: Document, markdown: str, content_width: Inches) -> None:
    """Walk the Markdown and emit Word blocks."""
    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    total = len(lines)

    while i < total:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code
        if _FENCE_RE.match(line):
            i += 1
            block: List[str] = []
            while i < total and not _FENCE_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            i += 1
            _render_code(document, block)
            continue

        # Horizontal rule — used between sections, so it becomes whitespace
        # rather than a literal line.
        if _RULE_RE.match(stripped):
            i += 1
            continue

        # Heading
        match = _HEADING_RE.match(stripped)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            # The document already has a title page; a stray H1 becomes H2.
            paragraph = document.add_heading("", level=min(max(level, 2), 4))
            _add_inline(
                paragraph,
                text,
                color=ACCENT if level <= 2 else SLATE,
            )
            for run in paragraph.runs:
                run.font.name = BODY_FONT
            _keep_with_next(paragraph)
            i += 1
            continue

        # Table
        if _is_table_line(line):
            block = []
            while i < total and (_is_table_line(lines[i]) or _TABLE_DIVIDER_RE.match(lines[i])):
                block.append(lines[i])
                i += 1
            _render_table(document, block, content_width)
            continue

        # Block quote
        if _QUOTE_RE.match(stripped):
            block = []
            while i < total and _QUOTE_RE.match(lines[i].strip()):
                block.append(_QUOTE_RE.match(lines[i].strip()).group(1))
                i += 1
            _render_quote(document, block)
            continue

        # Lists
        if _ULIST_RE.match(line) or _OLIST_RE.match(line):
            while i < total and (_ULIST_RE.match(lines[i]) or _OLIST_RE.match(lines[i])):
                ordered = _OLIST_RE.match(lines[i])
                text = (
                    ordered.group(2)
                    if ordered
                    else _ULIST_RE.match(lines[i]).group(1)
                )
                i += 1
                # Absorb continuation lines so wrapped list items stay one item.
                while (
                    i < total
                    and lines[i].strip()
                    and not _ULIST_RE.match(lines[i])
                    and not _OLIST_RE.match(lines[i])
                    and not _is_table_line(lines[i])
                    and not _HEADING_RE.match(lines[i].strip())
                    and lines[i].startswith((" ", "\t"))
                ):
                    text += " " + lines[i].strip()
                    i += 1

                paragraph = document.add_paragraph(
                    style="List Number" if ordered else "List Bullet"
                )
                paragraph.paragraph_format.space_after = Pt(2)
                _add_inline(paragraph, text)
            document.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # Paragraph — gather until a blank line or a block construct.
        block = []
        while i < total and lines[i].strip():
            candidate = lines[i]
            if (
                _is_table_line(candidate)
                or _HEADING_RE.match(candidate.strip())
                or _FENCE_RE.match(candidate)
                or _QUOTE_RE.match(candidate.strip())
                or _RULE_RE.match(candidate.strip())
                or _ULIST_RE.match(candidate)
                or _OLIST_RE.match(candidate)
            ):
                break
            block.append(candidate.strip())
            i += 1

        if block:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(7)
            _add_inline(paragraph, " ".join(block))


# ─── Public API ──────────────────────────────────────────────────────────


def _extract_title_metadata(markdown: str) -> Tuple[str, str, Dict[str, str], str]:
    """
    Peel the generated header block off the front of the document.

    The prompt service emits a title line, a two-column metadata table and a
    provenance quote.  Those belong on a title page rather than in the flow,
    so they are lifted out here and the remainder is rendered as the body.
    """
    lines = markdown.replace("\r\n", "\n").split("\n")
    title = ""
    subtitle = ""
    metadata: Dict[str, str] = {}
    body_start = 0

    for index, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        match = _HEADING_RE.match(stripped)
        if match and len(match.group(1)) == 1 and not title:
            heading = match.group(2).strip()
            if "—" in heading:
                title, _, subtitle = heading.partition("—")
                title, subtitle = title.strip(), subtitle.strip()
            else:
                title = heading
            body_start = index + 1
            continue

        if title and _is_table_line(stripped):
            cells = _split_row(stripped)
            if len(cells) == 2 and cells[0] and cells[0] != "---":
                metadata[re.sub(r"[*`]", "", cells[0])] = cells[1]
            body_start = index + 1
            continue

        if title and (_TABLE_DIVIDER_RE.match(stripped) or _QUOTE_RE.match(stripped)):
            body_start = index + 1
            continue

        if title and _RULE_RE.match(stripped):
            body_start = index + 1
            continue

        if title:
            break

    return title, subtitle, metadata, "\n".join(lines[body_start:])


def render_narrative_docx(
    markdown: str,
    title: Optional[str] = None,
    subtitle: Optional[str] = None,
    metadata: Optional[Dict[str, str]] = None,
    footer_label: Optional[str] = None,
) -> bytes:
    """
    Render a SCADA Design Narrative Markdown document to .docx bytes.

    Args:
        markdown: The assembled narrative, header block included.
        title: Overrides the title parsed from the leading H1.
        subtitle: Overrides the subtitle parsed from the leading H1.
        metadata: Overrides the title-page facts parsed from the header table.
        footer_label: Text shown beside the page number.

    Returns:
        The .docx file as bytes.
    """
    parsed_title, parsed_subtitle, parsed_meta, body = _extract_title_metadata(markdown)

    title = title or parsed_title or "SCADA Design Narrative"
    subtitle = subtitle if subtitle is not None else parsed_subtitle
    metadata = metadata if metadata is not None else parsed_meta

    document = Document()
    _configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    content_width = Inches(
        (section.page_width - section.left_margin - section.right_margin) / 914400
    )

    document.core_properties.title = f"SCADA Design Narrative — {title}"
    document.core_properties.category = "Design Basis"
    document.core_properties.comments = (
        "Generated from an RTAC configuration export. Every statement is derived "
        "from the exported configuration."
    )

    _add_footer(document, footer_label or title)
    _add_title_block(document, title, subtitle, metadata)
    render_markdown(document, body, content_width)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ─── CLI ─────────────────────────────────────────────────────────────────


def main() -> None:
    import argparse
    from pathlib import Path

    ap = argparse.ArgumentParser(
        description="Render a SCADA Design Narrative Markdown file to .docx"
    )
    ap.add_argument("markdown", help="Path to the assembled narrative Markdown")
    ap.add_argument("-o", "--output", required=True, help="Path to write the .docx")
    ap.add_argument("--title", help="Override the document title")
    args = ap.parse_args()

    source = Path(args.markdown).read_text(encoding="utf-8")
    payload = render_narrative_docx(source, title=args.title)
    Path(args.output).write_bytes(payload)
    print(f"Wrote {args.output} ({len(payload):,} bytes)")


if __name__ == "__main__":
    main()
